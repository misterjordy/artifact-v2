"""Build DOCX from section outputs using python-docx."""

from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def _add_cui_banners(doc: Any, cui_marking: str) -> None:
    """Add CUI header/footer to every section in the document."""
    for section in doc.sections:
        header = section.header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = cui_marking
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hp.style.font.size = Pt(10)
        hp.style.font.bold = True

        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.text = f"{cui_marking}\nDistribution Statement: Controlled"
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_docx(
    section_outputs: dict[str, str],
    sections: list[dict[str, Any]],
    template_name: str,
    classification: str = "UNCLASSIFIED",
) -> bytes:
    """Assemble a DOCX document from section outputs.

    Args:
        section_outputs: dict mapping section key to synthesized text.
        sections: ordered list of section definitions from template.
        template_name: name of the document template.
        classification: overall classification marking.

    Returns:
        DOCX file as bytes.
    """
    doc = Document()

    is_cui = "CUI" in classification.upper()

    if is_cui:
        cui_marking = classification if classification != "CUI" else "CUI"
        _add_cui_banners(doc, f"CONTROLLED // {cui_marking}")

    # Cover page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(template_name)
    run.bold = True
    run.font.size = Pt(24)

    if is_cui:
        marking_para = doc.add_paragraph()
        marking_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mrun = marking_para.add_run(classification)
        mrun.bold = True
        mrun.font.size = Pt(14)

    doc.add_page_break()  # type: ignore[no-untyped-call]  # python-docx has no type stubs

    for section_def in sections:
        key = section_def["key"]
        title = section_def["title"]
        text = section_outputs.get(key, "")

        doc.add_heading(title, level=1)
        if text:
            for paragraph in text.split("\n\n"):
                paragraph = paragraph.strip()
                if paragraph:
                    doc.add_paragraph(paragraph)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
